"""Preenchimento validado e atômico de templates DOCX."""

from __future__ import annotations

import re
from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


class DocumentGenerationError(RuntimeError):
    """Erro de geração que impede a publicação de um DOCX incompleto."""


class MissingPlaceholderValueError(DocumentGenerationError):
    """O template exige chaves que não foram fornecidas."""


class UnresolvedPlaceholderError(DocumentGenerationError):
    """O documento continuou contendo placeholders após o preenchimento."""


class DocumentGenerator:
    """Gera um DOCX somente quando todos os placeholders são resolvidos."""

    def __init__(self, template_path: str | Path):
        self.template_path = Path(template_path)

    def generate(self, data: dict, output_path: str | Path) -> None:
        output = Path(output_path)
        if not self.template_path.is_file():
            raise FileNotFoundError(f"Template não encontrado: {self.template_path}")

        doc = Document(self.template_path)
        required = self._collect_placeholders(doc)
        missing = sorted(required - set(data))
        if missing:
            raise MissingPlaceholderValueError(
                "Valores ausentes para os placeholders: " + ", ".join(missing)
            )

        for paragraph in self._iter_paragraphs(doc):
            self._replace_in_paragraph(paragraph, data)

        unresolved = sorted(self._collect_placeholders(doc))
        if unresolved:
            raise UnresolvedPlaceholderError(
                "Placeholders não resolvidos: " + ", ".join(unresolved)
            )

        self._ensure_table_of_contents(doc)

        output.parent.mkdir(parents=True, exist_ok=True)
        temporary_path: Path | None = None
        try:
            with NamedTemporaryFile(
                suffix=".docx",
                prefix=f".{output.stem}_",
                dir=output.parent,
                delete=False,
            ) as temporary:
                temporary_path = Path(temporary.name)
            doc.save(temporary_path)
            Document(temporary_path)
            temporary_path.replace(output)
        except Exception as error:
            if temporary_path is not None:
                temporary_path.unlink(missing_ok=True)
            if isinstance(error, DocumentGenerationError):
                raise
            raise DocumentGenerationError(
                f"Não foi possível gerar {output.name}: {error}"
            ) from error

    @classmethod
    def _collect_placeholders(cls, doc) -> set[str]:
        return {
            match.group(1).strip()
            for paragraph in cls._iter_paragraphs(doc)
            for match in PLACEHOLDER_PATTERN.finditer(paragraph.text)
        }

    @classmethod
    def _iter_paragraphs(cls, doc):
        yield from doc.paragraphs
        yield from cls._iter_table_paragraphs(doc.tables)
        for section in doc.sections:
            for container in (section.header, section.footer):
                yield from container.paragraphs
                yield from cls._iter_table_paragraphs(container.tables)

    @classmethod
    def _iter_table_paragraphs(cls, tables):
        seen_cells: set[object] = set()
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_xml = cell._tc
                    if cell_xml in seen_cells:
                        continue
                    seen_cells.add(cell_xml)
                    yield from cell.paragraphs
                    yield from cls._iter_table_paragraphs(cell.tables)

    def _replace_in_paragraph(self, paragraph, data: dict) -> None:
        full_text = "".join(run.text for run in paragraph.runs)
        keys = {
            match.group(1).strip()
            for match in PLACEHOLDER_PATTERN.finditer(full_text)
        }
        for key in keys:
            placeholder = re.compile(r"\{\{\s*" + re.escape(key) + r"\s*\}\}")
            while True:
                full_text = "".join(run.text for run in paragraph.runs)
                match = placeholder.search(full_text)
                if match is None:
                    break
                self._replace_range(
                    paragraph,
                    match.start(),
                    match.end(),
                    str(data[key]),
                )

    @staticmethod
    def _ensure_table_of_contents(doc) -> None:
        headings = [
            paragraph
            for paragraph in doc.paragraphs
            if paragraph.text.strip().casefold() == "sumário"
        ]
        if not headings:
            return

        instructions = doc.element.body.xpath(".//w:instrText")
        if any("TOC" in (instruction.text or "").upper() for instruction in instructions):
            return

        paragraphs = doc.paragraphs
        heading_index = next(
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph._p is headings[0]._p
        )
        next_content = None
        for paragraph in paragraphs[heading_index + 1 :]:
            if paragraph.text.strip():
                next_content = paragraph
                break
            protected_content = paragraph._p.xpath(
                ".//w:drawing | .//w:pict | .//w:fldChar | .//w:instrText"
            )
            if not protected_content:
                paragraph._element.getparent().remove(paragraph._element)
        if next_content is not None:
            next_content.paragraph_format.page_break_before = True

        toc_xml = OxmlElement("w:p")
        headings[0]._p.addnext(toc_xml)
        toc = Paragraph(toc_xml, headings[0]._parent)

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin.set(qn("w:dirty"), "true")
        toc.add_run()._r.append(begin)

        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = ' TOC \\o "1-2" \\h \\z \\u '
        toc.add_run()._r.append(instruction)

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        toc.add_run()._r.append(separate)
        toc.add_run("Atualize o sumário ao abrir o documento.")

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        toc.add_run()._r.append(end)

        settings = doc.settings._element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    @staticmethod
    def _replace_range(paragraph, start: int, end: int, value: str) -> None:
        runs = paragraph.runs
        offset = 0
        first_index = None
        last_index = None
        prefix = ""
        suffix = ""

        for index, run in enumerate(runs):
            run_start = offset
            run_end = offset + len(run.text)
            if first_index is None and run_end > start:
                first_index = index
                prefix = run.text[: start - run_start]
            if first_index is not None and run_end >= end:
                last_index = index
                suffix = run.text[end - run_start :]
                break
            offset = run_end

        if first_index is None or last_index is None:
            raise DocumentGenerationError("Não foi possível mapear um placeholder nos runs")

        runs[first_index].text = prefix + value + suffix
        for index in range(first_index + 1, last_index + 1):
            runs[index].text = ""
