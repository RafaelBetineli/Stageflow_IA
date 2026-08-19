import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from document_generator import DocumentGenerator, MissingPlaceholderValueError


PROJECT_ROOT = Path(__file__).resolve().parents[1]


class DocumentGeneratorTests(unittest.TestCase):
    def test_report_templates_start_references_on_a_new_page(self) -> None:
        templates = sorted(
            (PROJECT_ROOT / "templates").glob("*/modelo_relatorio_*.docx")
        )
        self.assertTrue(templates)
        for template in templates:
            document = Document(template)
            headings = [
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.strip().upper() == "REFERÊNCIAS BIBLIOGRÁFICAS"
            ]
            placeholders = [
                paragraph
                for paragraph in document.paragraphs
                if "{{REFERENCIAS}}" in paragraph.text
            ]
            with self.subTest(template=template.name):
                self.assertEqual(1, len(headings))
                self.assertTrue(headings[0].paragraph_format.page_break_before)
                self.assertTrue(headings[0].paragraph_format.keep_with_next)
                self.assertEqual(1, len(placeholders))
                self.assertEqual(WD_ALIGN_PARAGRAPH.LEFT, placeholders[0].alignment)
                self.assertEqual(0, placeholders[0].paragraph_format.left_indent)
                self.assertEqual(
                    0, placeholders[0].paragraph_format.first_line_indent
                )

    def test_replaces_body_table_header_footer_and_split_runs(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            template = root / "template.docx"
            output = root / "output.docx"
            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("Aluno: {{NOME_")
            paragraph.add_run("ALUNO}}")
            doc.add_table(rows=1, cols=1).cell(0, 0).text = "RA: {{RA_ALUNO}}"
            doc.sections[0].header.paragraphs[0].text = "{{CABECALHO}}"
            doc.sections[0].footer.paragraphs[0].text = "{{RODAPE}}"
            doc.save(template)

            DocumentGenerator(template).generate(
                {
                    "NOME_ALUNO": "Aluno Exemplo",
                    "RA_ALUNO": "000000",
                    "CABECALHO": "Cabeçalho",
                    "RODAPE": "Rodapé",
                },
                output,
            )

            generated = Document(output)
            self.assertEqual("Aluno: Aluno Exemplo", generated.paragraphs[0].text)
            self.assertEqual("RA: 000000", generated.tables[0].cell(0, 0).text)
            self.assertEqual("Cabeçalho", generated.sections[0].header.paragraphs[0].text)
            self.assertEqual("Rodapé", generated.sections[0].footer.paragraphs[0].text)

    def test_missing_value_does_not_replace_existing_output(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            template = root / "template.docx"
            output = root / "output.docx"
            doc = Document()
            doc.add_paragraph("{{OBRIGATORIO}}")
            doc.save(template)
            output.write_bytes(b"conteudo anterior")

            with self.assertRaises(MissingPlaceholderValueError):
                DocumentGenerator(template).generate({}, output)

            self.assertEqual(b"conteudo anterior", output.read_bytes())


if __name__ == "__main__":
    unittest.main()
